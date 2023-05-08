from  PyQt5 import QtWidgets, QtGui, QtCore
from  PyQt5.QtCore import QFileInfo
from PyQt5.QtWidgets import QApplication, QMainWindow, QFileDialog, QMessageBox
import docx2pdf
import docx
import re
import pandas as  pd #Установил ещё openpyxl, т.к. выдавал ошибку
from functools import partial
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from ui_sertificate import Ui_SertificateWindow
from sqlalchemy import create_engine, Table, Column, Integer, String, MetaData
from Classes.SendMail import SendMail
from settings import settings
class WindowSertificate(QtWidgets.QMainWindow):
    def __init__(self):
        super(WindowSertificate, self).__init__()
        self.ui = Ui_SertificateWindow()
        self.ui.setupUi(self)
        
        self.ui.pushButton.clicked.connect(self.sertificate_sample)
        self.ui.pushButton_2.clicked.connect(self.add_labels_to_sertificate)
        self.ui.pushButton_3.clicked.connect(self.send_sertificate_to_email)
        self.ui.pushButton_6.clicked.connect(self.delete_label_text_sertificate)
        self.ui.pushButton_7.clicked.connect(self.exit)

        self.__mail = SendMail(settings.mail_server,
                               settings.mail_port,
                               settings.mail_login,
                               settings.mail_password)
        
    def databased(self):
        engine = create_engine('sqlite:///diploma.db', echo=True)
        metadata = MetaData()
        mytable = Table('sertificate', metadata,
                        Column('id', Integer, primary_key=True),
                        Column('second_name', String),
                        Column('first_name', String),
                        Column('patronymic', String),
                        Column('email', String))
        self.connection = engine.connect()
        self.result = self.connection.execute(mytable.select())

    def add_labels_to_sertificate(self):
        self.databased()

        for participant in self.result:
            item = (participant[1], participant[2], participant[3], participant[4])
            dictor = self.create_data(item)
            file_choice_sertificate = self.get_url_to_sertificate() 
            doc = docx.Document(file_choice_sertificate)
            style = doc.styles['Normal']
            font = style.font
            font.size = docx.shared.Pt(14)
            font.name = 'Times New Roman'

            paragraph = None
            count_entry_sertificate = 0
            new_text_sertificate = None

            for paragraph in doc.paragraphs:
                new_text_sertificate, count_entry_sertificate = re.subn('|'.join(dictor.keys()), lambda match: dictor[match.group()], paragraph.text)
            if count_entry_sertificate > 0:
                paragraph.text = new_text_sertificate


            name_file_sertificate = ' '.join(item[:3])
            doc.save(f'{name_file_sertificate}.docx')

        self.connection.close()

        
        
    def exit(self):
        from main_window import MainWindow
        self.main = MainWindow()
        self.main.show()
        self.close()
        
    def sertificate_sample(self):
        self.file_dialog = QFileDialog(self, 'Выбрать файл', 'C://')
        self.file_dialog.setFileMode(QFileDialog.ExistingFile)
        self.file_dialog.setNameFilter("Microsoft Word (*.doc *.docx)")
        if self.file_dialog.exec_() == QFileDialog.Accepted:
            self.file_choice = self.file_dialog.selectedFiles()[0]
            if QFileInfo(self.file_choice).suffix() not in ["doc", "docx"]:
                QMessageBox.warning(self, "Ошибка", "Выбранный файл не является документом Word (.doc или .docx)")
                return 
            self.ui.label_3.setText(self.file_choice)


    def delete_label_text_sertificate(self):
        self.ui.label_3.setText('')

    
    def get_url_to_sertificate(self):
        return self.ui.label_3.text()


    def create_data(self,item: list) -> dict:  # Создание словаря с метками и их значениями
        self.full_name = f'{item[0]} {item[1]} {item[2]}'
        
        self.full_name_dictionary = {
                            '{{full_name}}': self.full_name,
                                        }
        return self.full_name_dictionary



    def create_mail(self, item):
        self.mail = item[3]
        return self.mail

    def send_sertificate_to_email(self):
        self.databased()
        messages = []
        for participant in self.result:
            item = (participant[1], participant[2], participant[3], participant[5])
            dictor = self.create_data(item)
            file_choice_sertificate = self.get_url_to_sertificate() 
            doc = docx.Document(file_choice_sertificate)
            style = doc.styles['Normal']
            font = style.font
            font.size = docx.shared.Pt(14)
            font.name = 'Times New Roman'

            paragraph = None
            count_entry_sertificate = 0
            new_text_sertificate = None

            for paragraph in doc.paragraphs:
                new_text_sertificate, count_entry_sertificate = re.subn('|'.join(dictor.keys()), lambda match: dictor[match.group()], paragraph.text)
            if count_entry_sertificate > 0:
                paragraph.text = new_text_sertificate


            name_file_sertificate = ' '.join(item[:4])
            doc.save(f'{name_file_sertificate}.docx')
            
    
            doc_pdf_path = f'{name_file_sertificate}.pdf'
            docx2pdf.convert(f'{name_file_sertificate}.docx', doc_pdf_path)

            mail_to_deliver = self.create_mail(item)
            subject = 'Диплом'
            text = 'Спасибо за участие в олимпиаде! Сертификат во вложении письма'
            files = [doc_pdf_path]


            msg = MIMEMultipart()
            msg['From'] = settings.mail_login
            msg['To'] = mail_to_deliver
            msg['Subject'] = subject
    
            msg.attach(MIMEText(text))
            if files:
                for file in files:
                    with open(file, 'rb') as file_user:
                        attach = MIMEApplication(file_user.read(), _subtype='pdf')
                        attach.add_header('Content-Disposition', 'attachment', filename=f'{name_file_sertificate}.pdf')

                        msg.attach(attach)

            messages.append(msg)
        
        self.__mail.send_message(messages)
        self.connection.close()
    
