from  PyQt5 import QtWidgets, QtGui, QtCore
from  PyQt5.QtCore import QFileInfo
from PyQt5.QtWidgets import QApplication, QMainWindow, QFileDialog, QMessageBox
import docx2pdf
import docx
import re
import pandas as  pd #Установил ещё openpyxl, т.к. выдавал ошибку
from functools import partial
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from ui_diploma import Ui_DiplomaWindow
from sqlalchemy import create_engine, Table, Column, Integer, String, MetaData
from Classes.SendMail import SendMail
from settings import settings
class WindowDiploma(QtWidgets.QMainWindow):
    def __init__(self):
        super(WindowDiploma, self).__init__()
        self.ui = Ui_DiplomaWindow()
        self.ui.setupUi(self)
        
        self.ui.pushButton.clicked.connect(self.diploma_sample)
        self.ui.pushButton_2.clicked.connect(self.add_labels_to_diploma)
        self.ui.pushButton_3.clicked.connect(self.send_diploma_to_email)
        self.ui.pushButton_6.clicked.connect(self.delete_label_text_diploma)
        self.ui.pushButton_7.clicked.connect(self.exit)

        self.__mail = SendMail(settings.mail_server,
                               settings.mail_port,
                               settings.mail_login,
                               settings.mail_password)
        
    def databased(self):
        engine = create_engine('sqlite:///diploma.db', echo=True)
        metadata = MetaData()
        mytable = Table('diploma', metadata,
                        Column('id', Integer, primary_key=True),
                        Column('last_name', String),
                        Column('first_name', String),
                        Column('patronymic', String),
                        Column('place', String),
                        Column('email', String))
        self.connection = engine.connect()
        self.result = self.connection.execute(mytable.select())

    def add_labels_to_diploma(self):
        self.databased()

        for participant in self.result:
            item = (participant[1], participant[2], participant[3], str(participant[4]), participant[5])
            dictor = self.create_data(item)
            file_choice_diploma = self.get_url_to_diploma() 
            doc = docx.Document(file_choice_diploma)
            style = doc.styles['Normal']
            font = style.font
            font.size = docx.shared.Pt(14)
            font.name = 'Times New Roman'

            paragraph = None
            count_entry_diploma = 0
            new_text_diploma = None

            for paragraph in doc.paragraphs:
                new_text_diploma, count_entry_diploma = re.subn('|'.join(dictor.keys()), lambda match: dictor[match.group()], paragraph.text)
            if count_entry_diploma > 0:
                paragraph.text = new_text_diploma


            name_file_diploma = ' '.join(item[:4])
            doc.save(f'{name_file_diploma} место.docx')

        self.connection.close()

        
        
    def exit(self):
        from main_window import MainWindow
        self.main = MainWindow()
        self.main.show()
        self.close()
        
    def diploma_sample(self):
        self.file_dialog = QFileDialog(self, 'Выбрать файл', 'C://')
        self.file_dialog.setFileMode(QFileDialog.ExistingFile)
        self.file_dialog.setNameFilter("Microsoft Word (*.doc *.docx)")
        if self.file_dialog.exec_() == QFileDialog.Accepted:
            self.file_choice = self.file_dialog.selectedFiles()[0]
            if QFileInfo(self.file_choice).suffix() not in ["doc", "docx"]:
                QMessageBox.warning(self, "Ошибка", "Выбранный файл не является документом Word (.doc или .docx)")
                return 
            self.ui.label_2.setText(self.file_choice)




    def delete_label_text_diploma(self):
        self.ui.label_2.setText('')


    def get_url_to_diploma(self):
        return self.ui.label_2.text()
  
    

    def create_data(self,item: list) -> dict:  # Создание словаря с метками и их значениями
        self.full_name = f'{item[0]} {item[1]} {item[2]}'
        self.place = item[3]
        self.full_name_dictionary_and_place = {
                            '{{full_name}}': self.full_name,
                            '{{place}}': self.place,     
                                        }
        return self.full_name_dictionary_and_place

    
    
    def create_mail(self, item):
        self.mail = item[4]
        return self.mail

    def send_diploma_to_email(self):
        self.databased()
        messages = []
        for participant in self.result:
            item = (participant[1], participant[2], participant[3], str(participant[4]), participant[5])
            dictor = self.create_data(item)
            file_choice_diploma = self.get_url_to_diploma() 
            doc = docx.Document(file_choice_diploma)
            style = doc.styles['Normal']
            font = style.font
            font.size = docx.shared.Pt(14)
            font.name = 'Times New Roman'

            paragraph = None
            count_entry_diploma = 0
            new_text_diploma = None

            for paragraph in doc.paragraphs:
                new_text_diploma, count_entry_diploma = re.subn('|'.join(dictor.keys()), lambda match: dictor[match.group()], paragraph.text)
            if count_entry_diploma > 0:
                paragraph.text = new_text_diploma


            name_file_diploma = ' '.join(item[:4])
            doc.save(f'{name_file_diploma} место.docx')
            
            # Convert the docx file to pdf
            doc_pdf_path = f'{name_file_diploma} место.pdf'
            docx2pdf.convert(f'{name_file_diploma} место.docx', doc_pdf_path)

            mail_to_deliver = self.create_mail(item)
            subject = 'Диплом'
            text = 'Вы выиграли в олимпиаде! Диплом во вложении письма'
            files = [doc_pdf_path]


            msg = MIMEMultipart()
            msg['From'] = settings.mail_login
            msg['To'] = mail_to_deliver
            msg['Subject'] = subject
            # тут текст из text
            msg.attach(MIMEText(text))
            if files:
                for file in files:
                    with open(file, 'rb') as file_user:
                        attach = MIMEApplication(file_user.read(), _subtype='pdf')
                        attach.add_header('Content-Disposition', 'attachment', filename=f'{name_file_diploma} место.pdf')

                        msg.attach(attach)

            messages.append(msg)
        
        self.__mail.send_message(messages)
        self.connection.close()
    
    

