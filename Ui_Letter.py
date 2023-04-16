from  PyQt5 import QtWidgets, QtGui, QtCore
from  PyQt5.QtCore import QFileInfo
from PyQt5.QtWidgets import QApplication, QMainWindow, QFileDialog, QMessageBox
import docx
import re
import pandas as  pd #Установил ещё openpyxl, т.к. выдавал ошибку
from functools import partial
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from letter import Ui_LetterWindow
class Window2(QtWidgets.QMainWindow):
    def __init__(self):
        super(Window2, self).__init__()
        self.ui = Ui_LetterWindow()
        self.ui.setupUi(self)
        self.ui.pushButton.clicked.connect(self.letter_sample)
        self.ui.pushButton_2.clicked.connect(self.databased)
        self.ui.pushButton_3.clicked.connect(self.add_labels_to_letter)
        self.ui.pushButton_4.clicked.connect(self.send_letter_to_email)
        self.ui.pushButton_6.clicked.connect(self.delete_label_text_letter)
        self.ui.pushButton_7.clicked.connect(self.delete_label_text_databased)
        
    def letter_sample(self):
        self.file_dialog = QFileDialog(self, 'Выбрать файл', 'C://')
        self.file_dialog.setFileMode(QFileDialog.ExistingFile)
        self.file_dialog.setNameFilter("Microsoft Word (*.doc *.docx)")
        if self.file_dialog.exec_() == QFileDialog.Accepted:
            self.file_choice = self.file_dialog.selectedFiles()[0]
            if QFileInfo(self.file_choice).suffix() not in ["doc", "docx"]:
                QMessageBox.warning(self, "Ошибка", "Выбранный файл не является документом Word (.doc или .docx)")
                return
            self.ui.label_2.setText(self.file_choice)

    def databased(self):
        self.file_dialog = QFileDialog(self, 'Выбрать файл', 'C://')
        self.file_dialog.setFileMode(QFileDialog.ExistingFile)
        self.file_dialog.setNameFilter("Microsoft Excel (*.xlsx *.xls)")
        if self.file_dialog.exec_() == QFileDialog.Accepted:
            self.file_choice = self.file_dialog.selectedFiles()[0]
            if QFileInfo(self.file_choice).suffix() not in ['xlsx', 'xls']:
                QMessageBox.warning(self, "Ошибка", "Выбранный файл не является документом Excel (.xlsx или .xls)")
                return
            self.ui.label_3.setText(self.file_choice)


    def delete_label_text_letter(self):
        self.ui.label_2.setText('')

    def delete_label_text_databased(self):
        self.ui.label_3.setText('')

    def get_url_to_letter(self):
        return self.ui.label_2.text(), self.ui.label_3.text()

    def create_data(self,item: list) -> dict:  # Создание словаря с метками и их значениями
        self.full_name = f'{item[0]} {item[1]} {item[2]}'
        self.place = item[3]
        self.full_name_dictionary_and_place = {
                            '{{full_name}}': self.full_name,
                            '{{place}}': self.place,     
                                        }
        return self.full_name_dictionary_and_place



    def add_labels_to_letter(self):
        self.file_choice_letter, self.file_choice_data = self.get_url_to_letter()
        self.letter_data = pd.read_excel(self.file_choice_data)
        self.data_list_letter = self.letter_data.values.tolist()
        for self.i, self.item in enumerate(self.data_list_letter):
            self.dictor = self.create_data(self.item)

            self.doc = docx.Document(self.file_choice_letter)
        
            self.style = self.doc.styles['Normal']
            self.font = self.style.font
            self.font.size = docx.shared.Pt(14)
            self.font.name = 'Times New Roman'

            for self.paragraph in self.doc.paragraphs:
                self.new_text_diploma, self.count_entry_diploma = re.subn('|'.join(self.dictor.keys()), lambda match: self.dictor[match.group()], self.paragraph.text)
            if self.count_entry_diploma > 0:
                self.paragraph.text = self.new_text_diploma
            self.name_file_letter = ' '.join(self.item[:3])
            self.doc.save(f'{self.name_file_letter}.docx')

            
    def create_mail(self, item):
        self.mail = item[4]
        return self.mail

    

    def send_letter_to_email(self):
        self.file_choice_letter, self.file_choice_data = self.get_url_to_letter()
        self.letter_data = pd.read_excel(self.file_choice_data)
        self.data_list_letter = self.letter_data.values.tolist()
        for self.i, self.item in enumerate(self.data_list_letter):
            self.dictor = self.create_data(self.item)

            self.doc = docx.Document(self.file_choice_letter)
        
            self.style = self.doc.styles['Normal']
            self.font = self.style.font
            self.font.size = docx.shared.Pt(14)
            self.font.name = 'Times New Roman'

            for self.paragraph in self.doc.paragraphs:
                self.new_text_diploma, self.count_entry_diploma = re.subn('|'.join(self.dictor.keys()), lambda match: self.dictor[match.group()], self.paragraph.text)
            if self.count_entry_diploma > 0:
                self.paragraph.text = self.new_text_diploma
            self.name_file_letter = ' '.join(self.item[:3])
            self.doc.save(f'{self.name_file_letter}.docx')

            self.mail_to_deliver = self.create_mail(self.item)
            self.subject = 'Диплом'
            self.text = 'Вы выиграли в олимпиаде! Диплом во вложении письма'
            self.files = [f'{self.name_file_letter}.docx']
            
            self.smtp_server = 'smtp.mail.ru'
            self.smtp_port = 587
            self.smtp_user = 'bezrukov30.00@mail.ru'
            self.smtp_password = ''#here was password (I delete it)

            self.server = smtplib.SMTP(self.smtp_server, self.smtp_port)
            self.server.starttls()
            self.server.login(self.smtp_user, self.smtp_password)
            
            self.msg = MIMEMultipart()
            self.msg['From'] = self.smtp_user
            self.msg['To'] = self.mail_to_deliver
            self.msg['Subject'] = self.subject

            self.msg.attach(MIMEText(self.text))
            if self.files:
                for self.file in self.files:
                    with open(self.file, 'rb') as self.file_user:
                        self.attach = MIMEApplication(self.file_user.read(), _subtype='docx')
                        self.attach.add_header('Content-Disposition', 'attachment', filename=self.file)
                        self.msg.attach(self.attach)
            self.server.sendmail(self.smtp_user, self.mail_to_deliver, self.msg.as_string())
            self.server.quit()
