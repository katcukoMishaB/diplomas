from  PyQt5 import QtWidgets, QtGui, QtCore
from  PyQt5.QtCore import QFileInfo
from PyQt5.QtWidgets import QApplication, QMainWindow, QFileDialog, QMessageBox
import sys
import docx
import re
import pandas as  pd #Установил ещё openpyxl, т.к. выдавал ошибку
from functools import partial
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication


#Выбор файлов
def diploma_sample():
    file_dialog = QFileDialog(window, 'Выбрать файл', 'C://')
    file_dialog.setFileMode(QFileDialog.ExistingFile)
    file_dialog.setNameFilter("Microsoft Word (*.doc *.docx)")
    if file_dialog.exec_() == QFileDialog.Accepted:
        file_choice = file_dialog.selectedFiles()[0]
        if QFileInfo(file_choice).suffix() not in ["doc", "docx"]:
            QMessageBox.warning(window, "Ошибка", "Выбранный файл не является документом Word (.doc или .docx)")
            return
        label.setText(file_choice)


def letter_sample():
    file_dialog = QFileDialog(window, 'Выбрать файл', 'C://')
    file_dialog.setFileMode(QFileDialog.ExistingFile)
    file_dialog.setNameFilter("Microsoft Word (*.doc *.docx)")
    if file_dialog.exec_() == QFileDialog.Accepted:
        file_choice = file_dialog.selectedFiles()[0]
        if QFileInfo(file_choice).suffix() not in ["doc", "docx"]:
            QMessageBox.warning(window, "Ошибка", "Выбранный файл не является документом Word (.doc или .docx)")
            return
        label3.setText(file_choice)

def databased():
    file_dialog = QFileDialog(window, 'Выбрать файл', 'C://')
    file_dialog.setFileMode(QFileDialog.ExistingFile)
    file_dialog.setNameFilter("Microsoft Excel (*.xlsx *.xls)")
    if file_dialog.exec_() == QFileDialog.Accepted:
        file_choice = file_dialog.selectedFiles()[0]
        if QFileInfo(file_choice).suffix() not in ['xlsx', 'xls']:
            QMessageBox.warning(window, "Ошибка", "Выбранный файл не является документом Excel (.xlsx или .xls)")
            return
        label2.setText(file_choice)

def delete_label_text_diploma():
    label.setText('')

def delete_label_text_letter():
    label3.setText('')

def delete_label_text_databased():
    label2.setText('')

def get_url_to_diploma():
    return label.text(), label2.text()

def get_url_to_letter():
    return label3.text(), label2.text()

def create_data(item: list) -> dict:  # Создание словаря с метками и их значениями
    full_name = f'{item[0]} {item[1]} {item[2]}'
    place = item[3]
    full_name_dictionary_and_place = {
                        '{{full_name}}': full_name,
                        '{{place}}': place,     
                                      }
    return full_name_dictionary_and_place



def add_labels_to_letter():
    file_choice_letter, file_choice_data = get_url_to_letter()
    letter_data = pd.read_excel(file_choice_data)
    data_list_letter = letter_data.values.tolist()
    for i, item in enumerate(data_list_letter):
        dictor = create_data(item)

        doc = docx.Document(file_choice_letter)
    
        style = doc.styles['Normal']
        font = style.font
        font.size = docx.shared.Pt(14)
        font.name = 'Times New Roman'

        for paragraph in doc.paragraphs:
            new_text_diploma, count_entry_diploma = re.subn('|'.join(dictor.keys()), lambda match: dictor[match.group()], paragraph.text)
        if count_entry_diploma > 0:
            paragraph.text = new_text_diploma
        name_file_letter = ' '.join(item[:3])
        doc.save(f'{name_file_letter}.docx')

def add_labels_to_diploma():
    #Считаваю данные из label'ов и передаю их в переменные в виде строки
    file_choice_diploma, file_choice_data = get_url_to_diploma()
    # Открытие выбранного файла с участниками и местами
    df = pd.read_excel(file_choice_data)
    data_list = df.values.tolist()  # преобразование всего датафрейма в список списков
    #Прохожу по всем вложенным спискам в списке
    for i, item in enumerate(data_list):
        # Создание словаря с метками и их значениями
        dictor = create_data(item)
        # Открытие выбранного документа 
        doc = docx.Document(file_choice_diploma)
        # Установка стилей документа
        style = doc.styles['Normal']
        font = style.font
        font.size = docx.shared.Pt(14)
        font.name = 'Times New Roman'
        # Замена меток на значения в документе

        for paragraph in doc.paragraphs:
            new_text_diploma, count_entry_diploma = re.subn('|'.join(dictor.keys()), lambda match: dictor[match.group()], paragraph.text)
        if count_entry_diploma > 0:
            paragraph.text = new_text_diploma

        # Сохранение документа с ФИО участника в качестве имени файла
        name_file_diploma = ' '.join(item[:4])
        doc.save(f'{name_file_diploma} место.docx')
        
def create_mail(item):
    mail = item[4]
    return mail

def send_diploma_to_email():
    #Считаваю данные из label'ов и передаю их в переменные в виде строки
    file_choice_diploma, file_choice_data = get_url_to_diploma()
    # Открытие выбранного файла с участниками и местами
    df = pd.read_excel(file_choice_data)
    data_list = df.values.tolist()  # преобразование всего датафрейма в список списков
    #Прохожу по всем вложенным спискам в списке
    for i, item in enumerate(data_list):
        # Создание словаря с метками и их значениями
        dictor = create_data(item)
        # Открытие выбранного документа 
        doc = docx.Document(file_choice_diploma)
        # Установка стилей документа
        style = doc.styles['Normal']
        font = style.font
        font.size = docx.shared.Pt(14)
        font.name = 'Times New Roman'
        # Замена меток на значения в документе

        for paragraph in doc.paragraphs:
            new_text_diploma, count_entry_diploma = re.subn('|'.join(dictor.keys()), lambda match: dictor[match.group()], paragraph.text)
        if count_entry_diploma > 0:
            paragraph.text = new_text_diploma

        # Сохранение документа с ФИО участника в качестве имени файла
        name_file_diploma = ' '.join(item[:4])
        doc.save(f'{name_file_diploma} место.docx')

        mail_to_deliver = create_mail(item)
        subject = 'Диплом'
        text = 'Вы выиграли в олимпиаде! Диплом во вложении письма'
        files = [f'{name_file_diploma} место.docx']
        # Параметры соединения с SMTP 
        smtp_server = 'smtp.mail.ru'
        smtp_port = 587
        smtp_user = 'bezrukov30.00@mail.ru'
        smtp_password = ''#here was password (I delete it)

        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(smtp_user, smtp_password)
        # Тут кому отправляем и от кого
        msg = MIMEMultipart()
        msg['From'] = smtp_user
        msg['To'] = mail_to_deliver
        msg['Subject'] = subject
        # тут текст из text
        msg.attach(MIMEText(text))
        # Тут добавляем файл ворд в письмо
        if files:
            for file in files:
                with open(file, 'rb') as file_user:
                    attach = MIMEApplication(file_user.read(), _subtype='docx')
                    attach.add_header('Content-Disposition', 'attachment', filename=file)
                    msg.attach(attach)

        server.sendmail(smtp_user, mail_to_deliver, msg.as_string())
        server.quit()

def send_letter_to_email():
    file_choice_letter, file_choice_data = get_url_to_letter()
    letter_data = pd.read_excel(file_choice_data)
    data_list_letter = letter_data.values.tolist()
    for i, item in enumerate(data_list_letter):
        dictor = create_data(item)

        doc = docx.Document(file_choice_letter)
    
        style = doc.styles['Normal']
        font = style.font
        font.size = docx.shared.Pt(14)
        font.name = 'Times New Roman'

        for paragraph in doc.paragraphs:
            new_text_diploma, count_entry_diploma = re.subn('|'.join(dictor.keys()), lambda match: dictor[match.group()], paragraph.text)
        if count_entry_diploma > 0:
            paragraph.text = new_text_diploma
        name_file_letter = ' '.join(item[:3])
        doc.save(f'{name_file_letter}.docx')

        mail_to_deliver = create_mail(item)
        subject = 'Диплом'
        text = 'Вы выиграли в олимпиаде! Диплом во вложении письма'
        files = [f'{name_file_letter}.docx']
         
        smtp_server = 'smtp.mail.ru'
        smtp_port = 587
        smtp_user = 'bezrukov30.00@mail.ru'
        smtp_password = ''#here was password (I delete it)

        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(smtp_user, smtp_password)
        
        msg = MIMEMultipart()
        msg['From'] = smtp_user
        msg['To'] = mail_to_deliver
        msg['Subject'] = subject

        msg.attach(MIMEText(text))
        if files:
            for file in files:
                with open(file, 'rb') as file_user:
                    attach = MIMEApplication(file_user.read(), _subtype='docx')
                    attach.add_header('Content-Disposition', 'attachment', filename=file)
                    msg.attach(attach)
        server.sendmail(smtp_user, mail_to_deliver, msg.as_string())
        server.quit()

app = QApplication(sys.argv)
window = QMainWindow()
window.setWindowTitle('Генератор дипломов и благодарственных писем')
window.setGeometry(500, 260, 700, 500)
# ------------------------------------------------------------------------
button_generator_diploma = QtWidgets.QPushButton(window)
button_generator_diploma.move(190, 350)
button_generator_diploma.setText('Генерация диплома')
button_generator_diploma.setFixedWidth(150)
button_generator_diploma.clicked.connect(add_labels_to_diploma)
# ------------------------------------------------------------------------
button_email_send_diploma = QtWidgets.QPushButton(window)
button_email_send_diploma.move(190, 450)
button_email_send_diploma.setText('Отправка диплома')
button_email_send_diploma.setFixedWidth(150)
button_email_send_diploma.clicked.connect(send_diploma_to_email)
# ------------------------------------------------------------------------
button_email_send_letter = QtWidgets.QPushButton(window)
button_email_send_letter.move(400, 450)
button_email_send_letter.setText('Отправка \n благодарственного письма\n ')
button_email_send_letter.setFixedWidth(150)
button_email_send_letter.clicked.connect(send_letter_to_email)
# ------------------------------------------------------------------------
button_generator_letter = QtWidgets.QPushButton(window)
button_generator_letter.move(400, 350)
button_generator_letter.setText('Генерация \n благодарственного письма')
button_generator_letter.setFixedWidth(150)
button_generator_letter.clicked.connect(add_labels_to_letter)
#------------------------------------------------------------------------
button_diploma = QtWidgets.QPushButton(window)
button_diploma.move(350, 150)
button_diploma.setText('Выбрать шаблон диплома')
button_diploma.setFixedWidth(160)
button_diploma.clicked.connect(diploma_sample)

label = QtWidgets.QLabel(window)
label.setFixedWidth(250)
label.move(30, 150)
# ------------------------------------------------------------------------
button_letter = QtWidgets.QPushButton(window)
button_letter.move(350, 40)
button_letter.setText('Выбрать шаблон \n благодарственного письма')
button_letter.setFixedWidth(160)
button_letter.clicked.connect(letter_sample)

label3 = QtWidgets.QLabel(window)
label3.setFixedWidth(250)
label3.move(30, 40)
# ------------------------------------------------------------------------
button_databased = QtWidgets.QPushButton(window)
button_databased.move(350, 250)
button_databased.setText('Выбрать данные')
button_databased.setFixedWidth(160)
button_databased.clicked.connect(databased)

label2 = QtWidgets.QLabel(window)
label2.setFixedWidth(250)
label2.move(30, 250)
# ------------------------------------------------------------------------
button_delete_databased = QtWidgets.QPushButton(window)
button_delete_databased.move(550, 250)
button_delete_databased.setText('Очистить')
button_delete_databased.setFixedWidth(100)
button_delete_databased.clicked.connect(delete_label_text_databased)
# ------------------------------------------------------------------------
button_delete_diploma= QtWidgets.QPushButton(window)
button_delete_diploma.move(550, 150)
button_delete_diploma.setText('Очистить')
button_delete_diploma.setFixedWidth(100)
button_delete_diploma.clicked.connect(delete_label_text_diploma)
# ------------------------------------------------------------------------
button_delete_letter = QtWidgets.QPushButton(window)
button_delete_letter.move(550, 40)
button_delete_letter.setText('Очистить')
button_delete_letter.setFixedWidth(100)
button_delete_letter.clicked.connect(delete_label_text_letter)
window.show()
sys.exit(app.exec())


